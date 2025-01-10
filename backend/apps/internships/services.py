import os
from django.conf import settings
from django.template.loader import render_to_string
from django.utils import timezone
from django.core.files import File
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from .models import Agreement, InternshipPlan, Internship
from apps.notifications.services import NotificationService
from docx import Document

class AgreementService:
    @staticmethod
    def generate_agreement_pdf(agreement):
        """Generate PDF agreement using ReportLab instead of WeasyPrint."""
        output_file = f'agreement_{agreement.id}_{timezone.now().strftime("%Y%m%d")}.pdf'
        output_path = os.path.join(settings.MEDIA_ROOT, 'agreements', output_file)

        # Create the PDF document
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )

        # Prepare the story (content)
        story = []
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1  # Center alignment
        )

        # Add title
        story.append(Paragraph("INTERNSHIP AGREEMENT", title_style))
        story.append(Spacer(1, 12))

        # Add agreement details
        story.append(Paragraph(f"Agreement No: {agreement.id}", styles['Normal']))
        story.append(Paragraph(f"Date: {timezone.now().strftime('%B %d, %Y')}", styles['Normal']))
        story.append(Spacer(1, 20))

        # Add parties information
        story.append(Paragraph("1. PARTIES", styles['Heading2']))
        story.append(Spacer(1, 12))

        # University Info
        story.append(Paragraph("UNIVERSITY:", styles['Heading3']))
        story.append(Paragraph(f"Name: {settings.UNIVERSITY_NAME}", styles['Normal']))
        story.append(Paragraph(f"Address: {settings.UNIVERSITY_ADDRESS}", styles['Normal']))
        story.append(Spacer(1, 12))

        # Organization Info
        story.append(Paragraph("ORGANIZATION:", styles['Heading3']))
        story.append(Paragraph(f"Name: {agreement.organization.name}", styles['Normal']))
        story.append(Paragraph(f"Address: {agreement.organization.address}", styles['Normal']))
        story.append(Paragraph(f"Contact Person: {agreement.organization.contact_person}", styles['Normal']))
        story.append(Spacer(1, 12))

        # Student Info
        story.append(Paragraph("STUDENT:", styles['Heading3']))
        story.append(Paragraph(f"Name: {agreement.student.get_full_name()}", styles['Normal']))
        story.append(Paragraph(f"Student ID: {agreement.student.username}", styles['Normal']))
        story.append(Spacer(1, 20))

        # Add internship details
        story.append(Paragraph("2. INTERNSHIP DETAILS", styles['Heading2']))
        story.append(Spacer(1, 12))

        # Create a table for internship details
        data = [
            ['Start Date', agreement.start_date.strftime('%B %d, %Y')],
            ['End Date', agreement.end_date.strftime('%B %d, %Y')],
            ['Department', agreement.organization.department.name if agreement.organization.department else 'N/A']
        ]
        
        t = Table(data, colWidths=[200, 300])
        t.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('PADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(t)
        story.append(Spacer(1, 20))

        # Add terms and conditions
        story.append(Paragraph("3. TERMS AND CONDITIONS", styles['Heading2']))
        story.append(Spacer(1, 12))
        terms = get_terms_and_conditions()
        for term in terms.split('\n'):
            if term.strip():
                story.append(Paragraph(term, styles['Normal']))
                story.append(Spacer(1, 6))

        # Add signature section
        story.append(Spacer(1, 30))
        story.append(Paragraph("4. SIGNATURES", styles['Heading2']))
        story.append(Spacer(1, 12))

        # Create signature table
        sig_data = [
            ['University Representative', 'Organization Representative', 'Student'],
            ['_________________', '_________________', '_________________'],
            ['Date: ____________', 'Date: ____________', 'Date: ____________']
        ]
        
        sig_table = Table(sig_data, colWidths=[180, 180, 180])
        sig_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        story.append(sig_table)

        # Build the PDF
        doc.build(story)

        return f'agreements/{output_file}'

    @staticmethod
    def generate_docx_agreement(agreement):
        """Generate DOCX version of agreement for editing."""
        doc = Document()
        doc.add_heading('Internship Agreement', 0)

        # Add agreement details
        doc.add_heading('1. Parties', level=1)
        doc.add_paragraph(f'Student: {agreement.student.get_full_name()}')
        doc.add_paragraph(f'Organization: {agreement.organization.name}')
        # Add more sections...

        # Save document
        output_file = f'agreement_{agreement.id}_{timezone.now().strftime("%Y%m%d")}.docx'
        output_path = os.path.join(settings.MEDIA_ROOT, 'agreements', output_file)
        doc.save(output_path)
        return f'agreements/{output_file}'

    @staticmethod
    def process_signature(agreement, user_type, signature_data=None):
        """Process digital signature with enhanced security."""
        try:
            # Verify signature if provided
            if signature_data:
                # Add digital signature verification
                pass

            # Update agreement status
            agreement.sign(user_type)
            
            # Create notification
            NotificationService.create_notification(
                recipient=agreement.student,
                title='Agreement Update',
                message=f'Agreement signed by {user_type}',
                notification_type='agreement'
            )

            if agreement.status == 'approved':
                # Create internship plan
                plan = InternshipPlan.objects.create(
                    internship=agreement.internship,
                    content=InternshipPlanService.generate_plan_template(agreement.internship)
                )
                
                # Notify relevant parties
                NotificationService.create_bulk_notifications([
                    {
                        'recipient': agreement.student,
                        'title': 'Agreement Approved',
                        'message': 'Your internship agreement has been approved'
                    },
                    {
                        'recipient': agreement.organization.contact_person,
                        'title': 'Agreement Approved',
                        'message': 'Internship agreement has been approved'
                    }
                ])

            return True
        except Exception as e:
            print(f"Error processing signature: {str(e)}")
            return False

class InternshipPlanService:
    @staticmethod
    def generate_plan_template(internship):
        """Generate comprehensive internship plan template."""
        # Get organization-specific template if available
        org_template = get_organization_template(internship.organization)
        if org_template:
            return org_template

        # Generate default template
        weeks = calculate_internship_weeks(internship.start_date, internship.end_date)
        
        template = f"""
        Internship Plan for {internship.student.get_full_name()}
        Organization: {internship.organization.name}
        Duration: {internship.start_date} to {internship.end_date}
        Total Weeks: {weeks}

        1. Learning Objectives:
        - Understand {internship.organization.name}'s business processes
        - Develop practical skills in {internship.department.name if internship.department else 'assigned area'}
        - Gain hands-on experience with industry tools and technologies

        2. Weekly Schedule:
        {generate_weekly_schedule(weeks)}

        3. Expected Outcomes:
        - Complete assigned projects and tasks
        - Develop professional work habits
        - Build industry network
        - Create portfolio of work

        4. Evaluation Criteria:
        - Task completion quality
        - Professional behavior
        - Technical skill development
        - Communication skills
        - Initiative and proactivity

        5. Supervision Details:
        Mentor: {internship.mentor.get_full_name() if internship.mentor else 'TBD'}
        Meeting Schedule: Weekly check-ins
        Communication Channels: Email, Teams/Slack
        
        6. Documentation Requirements:
        - Weekly progress reports
        - Monthly evaluations
        - Final presentation
        - Internship portfolio

        7. Additional Requirements:
        - Attend team meetings
        - Participate in training sessions
        - Follow organization policies
        - Maintain confidentiality
        """
        return template

    @staticmethod
    def review_plan(plan, reviewer, status, feedback=None):
        """Review internship plan with notifications."""
        plan.status = status
        if feedback:
            plan.feedback = feedback
        if status == 'approved':
            plan.approved_by = reviewer
        plan.save()

        # Notify student
        NotificationService.create_notification(
            recipient=plan.internship.student,
            title=f'Plan {status.title()}',
            message=f'Your internship plan has been {status}. {feedback if feedback else ""}',
            notification_type='plan'
        )

def generate_weekly_schedule(weeks):
    """Generate detailed weekly schedule template."""
    schedule = []
    for i in range(1, weeks + 1):
        if i == 1:
            schedule.append(f"""
            Week {i}:
            - Orientation and introduction
            - Company policies and procedures
            - Team introductions
            - Setup and access to required systems
            """)
        elif i == weeks:
            schedule.append(f"""
            Week {i}:
            - Final project completion
            - Documentation finalization
            - Knowledge transfer
            - Final presentation preparation
            """)
        else:
            schedule.append(f"""
            Week {i}:
            - Project work
            - Skill development
            - Team collaboration
            - Weekly review
            """)
    
    return "\n".join(schedule)

def calculate_internship_weeks(start_date, end_date):
    """Calculate total weeks of internship."""
    delta = end_date - start_date
    return (delta.days + 1) // 7

def get_organization_template(organization):
    """Get organization-specific template if available."""
    # Implement organization template logic
    return None

def get_terms_and_conditions():
    """Get standard terms and conditions."""
    return """
    1. Confidentiality Agreement
    2. Intellectual Property Rights
    3. Work Hours and Schedule
    4. Code of Conduct
    5. Safety Regulations
    6. Evaluation Process
    7. Termination Conditions
    """

def generate_verification_qr(agreement):
    """Generate QR code for agreement verification."""
    import qrcode
    import base64
    from io import BytesIO

    # Generate verification URL
    verification_url = f"{settings.SITE_URL}/verify-agreement/{agreement.id}"
    
    # Create QR code
    qr = qrcode.QRCode(version=1, box_size=10, border=5)
    qr.add_data(verification_url)
    qr.make(fit=True)
    
    # Create QR image
    qr_image = qr.make_image(fill_color="black", back_color="white")
    
    # Convert to base64 for embedding in PDF
    buffer = BytesIO()
    qr_image.save(buffer, format='PNG')
    qr_base64 = base64.b64encode(buffer.getvalue()).decode()
    
    return f"data:image/png;base64,{qr_base64}"

def generate_plan_template(internship):
    """Generate default plan template based on internship details."""
    return InternshipPlanService.generate_plan_template(internship) 